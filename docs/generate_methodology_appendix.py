#!/usr/bin/env python3
"""
Generate the AI Labor Forecast Methodology Appendix as a Word document (.docx).
Client-facing methodology bible — clear, sharp, in-depth.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "AI_Labor_Forecast_Methodology_Appendix.docx")

# Colors
NAVY = RGBColor(15, 30, 62)
SLATE = RGBColor(55, 65, 81)
ACCENT = RGBColor(37, 99, 235)
BODY_COLOR = RGBColor(38, 38, 42)
LIGHT_GRAY = RGBColor(148, 155, 168)


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "0F1E3E")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = BODY_COLOR
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F7FB")

    # Apply column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def build_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # -- Heading styles --
    for level, (size, color, spacing) in {
        1: (22, NAVY, 18),
        2: (14, NAVY, 12),
        3: (11, SLATE, 8),
    }.items():
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(spacing)
        h.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Labor Forecast")
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Methodology Appendix")
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    doc.add_paragraph("")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("_" * 30)
    run.font.color.rgb = ACCENT

    doc.add_paragraph("")

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("18-Month AI Labor Displacement Forecast\n2×2 Scenario Framework  ·  21 Industries  ·  462 Occupations")
    run.font.size = Pt(12)
    run.font.color.rgb = SLATE

    for _ in range(6):
        doc.add_paragraph("")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("March 2026  ·  Confidential")
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (manual)
    # =========================================================================
    doc.add_heading("Contents", level=1)
    toc_items = [
        ("1", "Executive Summary"),
        ("2", "The Master Equation"),
        ("3", "Three Analytical Layers"),
        ("4", "Layer 1 — Task-Level Automatability"),
        ("5", "Layer 2 — Job-Level Automatability"),
        ("6", "Layer 3 — Industry-Level Frictions"),
        ("7", "The 2×2 Scenario Framework"),
        ("8", "Validation & Quality Assurance"),
        ("9", "Key Design Principles"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}.   {title}")
        run.font.size = Pt(10.5)
        run.font.color.rgb = SLATE

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading("1.  Executive Summary", level=1)

    doc.add_paragraph(
        "The AI Labor Forecast is a bottom-up, scenario-driven model that estimates "
        "the share of employment at risk of displacement by AI within an 18-month horizon "
        "(through October 2027). It spans 21 industry sectors, 462 occupations, and over "
        "5,300 individual work tasks."
    )
    doc.add_paragraph(
        "Rather than making a single-point prediction, the model produces four displacement "
        "estimates per sector by crossing two independent axes of uncertainty:"
    )

    bullets = [
        ("Technology capability", "How far will frontier AI models advance — moderate, "
         "steady-state improvements, or a significant step-change with architectural breakthroughs?"),
        ("Adoption friction", "How quickly will organizations deploy AI at scale — encountering "
         "low institutional resistance, or high friction from legacy systems, regulation, and cultural inertia?"),
    ]
    for bold_part, rest in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bold_part + ". ")
        run.bold = True
        p.add_run(rest)

    doc.add_paragraph(
        "The result is a displacement rate for each sector×scenario combination, expressed as the "
        "percentage of current employment at risk over 18 months. This is not a prediction of layoffs — "
        "it is the upper envelope of roles where AI could economically substitute for human labor, "
        "before firms make strategic decisions about redeployment, reskilling, or growth."
    )

    doc.add_paragraph(
        "The model is multiplicative by design: displacement only registers when a job is technically "
        "automatable AND the workflow allows separation AND the industry adopts the technology AND "
        "regulation permits it. Any single bottleneck — a tightly coupled workflow, a liability regime, "
        "a slow-adopting sector — compresses the estimate toward zero."
    )

    doc.add_page_break()

    # =========================================================================
    # 2. THE MASTER EQUATION
    # =========================================================================
    doc.add_heading("2.  The Master Equation", level=1)

    doc.add_paragraph(
        "Every displacement estimate in the forecast is produced by one equation:"
    )

    # Equation block
    eq_p = doc.add_paragraph()
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_p.paragraph_format.space_before = Pt(12)
    eq_p.paragraph_format.space_after = Pt(12)
    run = eq_p.add_run("d(t; i, s)  =  d_max(i)  ×  S(a(i,s))  ×  E(i)  ×  T(t; i, s)  ×  R(i)")
    run.font.name = "Consolas"
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph("Where:")

    terms = [
        ("d(t; i, s)", "Displacement rate for industry i, scenario s, at time t (18 months)"),
        ("d_max(i)", "Displacement ceiling — the maximum rate at which jobs can structurally disappear in sector i, "
         "derived from historical labor turnover data"),
        ("S(a(i,s))", "Sigmoid-transformed automatability — how much of the sector's work AI can technically perform, "
         "compressed through a sigmoid to dampen extreme values"),
        ("E(i)", "Elasticity dampener — captures the Jevons paradox: does cheaper AI-driven output create new demand "
         "that absorbs displaced workers, or is demand fixed?"),
        ("T(t; i, s)", "Adoption velocity — a logistic S-curve capturing how fast the industry actually deploys AI, "
         "given both technological readiness and institutional frictions"),
        ("R(i)", "Structural resistance — hard regulatory floors, licensure mandates, and liability regimes that "
         "cap displacement regardless of technical capability"),
    ]
    for term, desc in terms:
        p = doc.add_paragraph()
        run = p.add_run(term)
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.add_run(f"  —  {desc}")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The multiplicative structure is the core design choice. Each factor acts as a gate: "
        "if any one of them is near zero, the product collapses. A job can be highly automatable "
        "(high S(a)) but still show minimal displacement if the industry is slow to adopt (low T), "
        "demand absorbs the productivity gains (low E), or regulation blocks deployment (low R). "
        "This prevents the model from generating implausible displacement estimates driven by a "
        "single dominant factor."
    )

    doc.add_page_break()

    # =========================================================================
    # 3. THREE ANALYTICAL LAYERS
    # =========================================================================
    doc.add_heading("3.  Three Analytical Layers", level=1)

    doc.add_paragraph(
        "The model builds displacement estimates from the bottom up through three nested layers:"
    )

    add_styled_table(doc,
        ["Layer", "Unit of Analysis", "Count", "Key Output"],
        [
            ["Task-Level", "Individual work tasks", "5,383 tasks × 2 scenarios", "Autonomy fraction per task"],
            ["Job-Level", "SOC occupations", "462 occupations", "Automatability score a(j,s)"],
            ["Industry-Level", "Economic sectors", "21 sectors × 4 scenarios", "Displacement rate d(t; i, s)"],
        ],
        col_widths=[3.0, 4.0, 4.5, 4.5],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Each layer feeds into the next. Task scores aggregate into job-level automatability; "
        "job scores combine with industry-level frictions to produce displacement estimates. "
        "This bottom-up architecture ensures that the macro forecast is grounded in the specific "
        "work content of each occupation, not broad assumptions about industry vulnerability."
    )

    doc.add_page_break()

    # =========================================================================
    # 4. LAYER 1 — TASK-LEVEL AUTOMATABILITY
    # =========================================================================
    doc.add_heading("4.  Layer 1 — Task-Level Automatability", level=1)

    doc.add_paragraph(
        "The foundation of the model is a task-by-task assessment of AI capability. Each of the "
        "5,383 tasks in the O*NET database is scored on a single question:"
    )

    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_before = Pt(8)
    q.paragraph_format.space_after = Pt(8)
    run = q.add_run(
        '"What fraction of this task\'s value-add can a frontier AI model deliver\n'
        'autonomously, at required quality and economic viability, by October 2027?"'
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = SLATE

    doc.add_heading("Autonomy Fraction Scale", level=2)

    add_styled_table(doc,
        ["Score", "Label", "Interpretation"],
        [
            ["0.00", "AI cannot perform", "Physical tasks, genuine human-only relationships — no path to automation"],
            ["0.25", "AI assists", "AI handles routine subtasks; human does core judgment (e.g., AI drafts summaries, human interprets)"],
            ["0.50", "Shared autonomy", "AI autonomously handles ~half of cases; complex cases need human oversight"],
            ["0.75", "AI leads", "AI performs in most cases; human needed for rare edge cases or final sign-off"],
            ["1.00", "Fully autonomous", "AI fully autonomous, no human involvement needed across essentially all cases"],
        ],
        col_widths=[1.5, 3.0, 11.5],
    )

    doc.add_heading("Two Technology Scenarios", level=2)

    doc.add_paragraph(
        "Every task is scored twice — once under each technology scenario. The two scenarios "
        "bound the plausible range of frontier AI capability by October 2027:"
    )

    doc.add_heading("Moderate Capability Gains", level=3)
    doc.add_paragraph(
        "Steady compounding of current approaches with no architectural breakthroughs. Context windows "
        "reach ~1.5M tokens (reliable to ~400K). AI agents operate in 4–6 hour autonomous blocks on "
        "5–8 step workflows. Tool use improves for familiar APIs. Token costs drop ~3×. The model "
        "often doesn't know when it's wrong."
    )

    doc.add_heading("Significant Capability Gains", level=3)
    doc.add_paragraph(
        "One or more step-change breakthroughs beyond the current trajectory. Context windows reach "
        "~2–3M tokens (reliable to ~1M). Agents operate in 6–10 hour blocks on 12–15 step workflows "
        "across unfamiliar systems. Reinforcement learning expands to semi-verifiable professional "
        "domains (finance, law, medicine). Parallel agent coordination becomes reliable. Token costs "
        "drop ~10×. Better error detection and more consistent outputs."
    )

    doc.add_heading("Universal Limitations", level=3)
    doc.add_paragraph(
        "Certain AI limitations bind equally in both scenarios, meaning tasks constrained by them "
        "receive identical scores regardless of the technology trajectory:"
    )
    limitations = [
        "No physical-world embodiment (all physical tasks score 0.00)",
        "No self-recursive model improvement or continual learning from deployment",
        "No accurate uncertainty calibration or robust causal reasoning over novel mechanisms",
        "No genuine theory of mind in adversarial settings",
        "Self-correction has a ceiling — catches surface errors, not subtle framing errors",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    doc.add_heading("Task Scoring Pipeline", level=2)
    doc.add_paragraph(
        "Task-level scoring is performed by a multi-agent pipeline using Claude Opus 4.6 at "
        "temperature 0.2, with structured quality controls:"
    )
    steps = [
        ("Calibration Agent", "Scores ~15 anchor tasks spanning the full difficulty range, "
         "with detailed chain-of-thought reasoning. These anchors serve as alignment exemplars."),
        ("Batch Scoring Agent", "Processes tasks in batches of 25–30, with calibration anchors "
         "embedded in every prompt. Each task receives an autonomy score for both scenarios."),
        ("Drift Auditor Agent", "Reviews all scores for internal inconsistencies, compares against "
         "calibration anchors, and re-scores any flagged tasks to maintain coherence."),
    ]
    for name, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}:  ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 5. LAYER 2 — JOB-LEVEL AUTOMATABILITY
    # =========================================================================
    doc.add_heading("5.  Layer 2 — Job-Level Automatability", level=1)

    doc.add_paragraph(
        "Task scores aggregate into a single automatability score per occupation. This is "
        "not a simple average — it accounts for the structure of work within each job."
    )

    doc.add_heading("The Job-Level Formula", level=2)

    eq2 = doc.add_paragraph()
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq2.paragraph_format.space_before = Pt(10)
    eq2.paragraph_format.space_after = Pt(10)
    run = eq2.add_run("a(j, s)  =  tc_adj(j, s)  ×  w(j)")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph("Then transformed through a sigmoid:")

    eq3 = doc.add_paragraph()
    eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq3.paragraph_format.space_before = Pt(8)
    eq3.paragraph_format.space_after = Pt(8)
    run = eq3.add_run("S(a)  =  a^k  /  (a^k + (1-a)^k)        where k = 0.8")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_heading("Component 1: Task-Coverage Depth, Adjusted (tc_adj) — Scenario-Dependent", level=2)

    doc.add_paragraph(
        "This component is computed programmatically from the task-level autonomy scores. "
        "It measures how deeply AI can penetrate a job's task portfolio:"
    )

    eq4 = doc.add_paragraph()
    eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq4.paragraph_format.space_before = Pt(8)
    eq4.paragraph_format.space_after = Pt(8)
    run = eq4.add_run("tc_adj  =  tc_mean  ×  (1 - z)^0.5  ×  min(1, h / tc)")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    tc_terms = [
        ("tc_mean", "Importance-weighted average of task autonomy scores (weighted by Time_Share_Pct)"),
        ("z", "Time share of zero-autonomy tasks. The square-root penalty (1-z)^0.5 means a job with 40% "
         "non-automatable work is penalized, but not zeroed out"),
        ("h / tc", "Ratio of high-autonomy task share (scores ≥ 0.65) to overall task coverage. "
         "Caps inflated scores from jobs with many moderately-automatable but few deeply-automatable tasks"),
    ]
    for term, desc in tc_terms:
        p = doc.add_paragraph()
        run = p.add_run(term)
        run.bold = True
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.add_run(f"  —  {desc}")

    doc.add_paragraph(
        "This adjusted metric prevents a common pitfall: a job where every task is 30% automatable "
        "should not score the same as a job where half the tasks are 60% automatable. The h/tc cap "
        "distinguishes shallow breadth from deep automation potential."
    )

    doc.add_heading("Component 2: Workflow Separability (w) — Scenario-Independent", level=2)

    doc.add_paragraph(
        "Even if many of a job's tasks are technically automatable, displacement requires that "
        "those tasks can be separated from the rest of the workflow. Workflow separability captures "
        "this structural property — it is a feature of how work is organized, not of AI capability, "
        "and therefore does not change across scenarios."
    )

    add_styled_table(doc,
        ["Score", "Label", "Description", "Exemplars"],
        [
            ["1.00", "Fully separable", "Independent tasks; can offload without disruption",
             "Data entry keyers, court reporters, financial clerks"],
            ["0.75", "Mostly separable", "Natural handoff boundaries; minimal coordination",
             "Paralegals, insurance underwriters, technical writers"],
            ["0.50", "Partially separable", "Tasks interleaved throughout workday; redesign needed",
             "Nurses, general managers, architects"],
            ["0.25", "Minimally separable", "Continuous flow; cannot disaggregate",
             "Physicians, psychologists, clergy, dentists"],
        ],
        col_widths=[1.5, 3.0, 5.5, 6.0],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "Workflow separability is scored by LLM agents informed by structured job profiles that "
        "include task distributions, work activity patterns, and workflow architecture indicators. "
        "The diagnostic questions are: Do automatable tasks cluster in time? What is the handoff "
        "friction? How tightly coupled is information flow between tasks?"
    )

    doc.add_heading("The Sigmoid Transform", level=2)

    doc.add_paragraph(
        "The raw automatability score a = tc_adj × w is passed through a sigmoid with k = 0.8. "
        "This is deliberately softer than a standard logistic — it preserves meaningful signal "
        "across a wider range while compressing extreme values:"
    )

    add_styled_table(doc,
        ["Raw Score (a)", "Transformed S(a)", "Effect"],
        [
            ["0.10", "0.07", "Low automatability compressed further"],
            ["0.30", "0.26", "Modest reduction"],
            ["0.50", "0.50", "Midpoint unchanged"],
            ["0.70", "0.74", "Modest boost"],
            ["0.90", "0.93", "High automatability compressed slightly"],
        ],
        col_widths=[3.5, 3.5, 9.0],
    )

    doc.add_heading("Job-Level Scoring Pipeline", level=2)

    phases = [
        ("Phase 0 — Profile Extraction",
         "Structured occupation profiles are extracted from the workbook: SOC code, title, sector, employment, "
         "wage, task statistics, Generalized Work Activity distributions, and construct-specific features "
         "(interpersonal share, digital share, judgment share, task heterogeneity)."),
        ("Phase 1 — Calibration",
         "30 boundary-case occupations are scored with full chain-of-thought reasoning to produce anchor "
         "exemplars. These span the full range of each variable and must complete before batch scoring begins."),
        ("Phase 2 — Batch Scoring",
         "462 occupations are scored in batches with calibration anchors embedded in every prompt. "
         "Scoring is parallelized across 6 sector-grouped agents. Each agent produces independent "
         "reasoning blocks per variable to mitigate halo effects."),
        ("Phase 3 — Audit",
         "Programmatic checks flag distribution skew (>35% at one anchor), high cross-variable "
         "correlation (>0.80), SOC-group incoherence (within-group σ > 0.40), and logical inconsistencies. "
         "Flagged jobs (~10–15%) are re-scored by an LLM auditor at temperature 0.1."),
        ("Phase 4 — Reliability Verification",
         "A stratified 10% sample (46 occupations) is independently re-scored. Agreement is measured via "
         "weighted quadratic kappa (threshold: κ ≥ 0.75), within-one-step agreement (≥ 90%), and exact "
         "agreement (≥ 65%)."),
        ("Phase 5 — Write-back",
         "Final validated scores are written to the master workbook for downstream computation."),
    ]
    for phase_name, desc in phases:
        p = doc.add_paragraph()
        run = p.add_run(f"{phase_name}.  ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 6. LAYER 3 — INDUSTRY-LEVEL FRICTIONS
    # =========================================================================
    doc.add_heading("6.  Layer 3 — Industry-Level Frictions", level=1)

    doc.add_paragraph(
        "Job-level automatability tells us what AI can do. Industry-level variables determine "
        "what actually happens — how fast, how much, and against what barriers. Four industry "
        "variables gate the final displacement estimate."
    )

    # -- d_max --
    doc.add_heading("d_max — Displacement Ceiling", level=2)

    p = doc.add_paragraph()
    run = p.add_run("What it captures: ")
    run.bold = True
    p.add_run(
        "The maximum rate at which jobs can structurally disappear in a sector within "
        "18 months, regardless of the cause. No industry has ever shed more than ~12% of "
        "employment per year, even during massive technological disruptions (typesetters, "
        "switchboard operators). d_max enforces this historical ceiling."
    )

    p = doc.add_paragraph()
    run = p.add_run("How it's derived: ")
    run.bold = True
    p.add_run(
        "From BLS JOLTS total separation rate (TSR) data — the monthly rate at which workers "
        "leave jobs through all channels (quits, layoffs, retirements). The monthly rate is "
        "compounded to an 18-month window:"
    )

    eq5 = doc.add_paragraph()
    eq5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq5.paragraph_format.space_before = Pt(8)
    eq5.paragraph_format.space_after = Pt(8)
    run = eq5.add_run("d_max  =  1 - (1 - monthly_rate / 100)^18")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Scenario-stable: ")
    run.bold = True
    p.add_run(
        "d_max does not vary by scenario. The structural speed at which an industry's workforce "
        "can turn over is determined by labor market dynamics (contracts, notice periods, hiring "
        "cycles), not by AI capability."
    )

    doc.add_paragraph("")

    # d_max table — selected sectors
    add_styled_table(doc,
        ["Sector", "Monthly TSR (%)", "d_max (18-mo)", "Intuition"],
        [
            ["Government & Public Admin", "1.38", "0.22", "Civil service protections, slow turnover"],
            ["Insurance", "1.70", "0.27", "Regulated, credentialed, stable workforce"],
            ["Law Firms", "1.80", "0.28", "Partnership track, high retention"],
            ["Finance & Banking", "2.15", "0.32", "Moderate turnover, regulated"],
            ["Technology & Software", "3.30", "0.45", "High voluntary churn, startup culture"],
            ["Retail Trade", "3.81", "0.50", "High turnover, low switching costs"],
            ["Accommodation & Food", "5.45", "0.64", "Highest churn sector in the economy"],
        ],
        col_widths=[4.5, 3.0, 3.0, 5.5],
    )

    doc.add_paragraph("")

    # -- E --
    doc.add_heading("E — Elasticity Dampener (Jevons Paradox)", level=2)

    p = doc.add_paragraph()
    run = p.add_run("What it captures: ")
    run.bold = True
    p.add_run(
        "When AI makes output cheaper, does demand expand to absorb displaced workers, or is "
        "demand fixed? This is the Jevons paradox applied to labor markets."
    )

    add_styled_table(doc,
        ["E Value", "Label", "Mechanism", "Examples"],
        [
            ["0.25", "Demand-absorbing", "Strong Jevons effect — cheaper output creates new demand "
             "that absorbs workers", "Technology (cheaper dev → more software), Healthcare diagnostics "
             "(cheaper screening → more screening)"],
            ["0.50", "Partially absorbing", "Mixed response — some new demand, but not enough to "
             "fully offset displacement", "Finance, Legal, Manufacturing, Retail"],
            ["1.00", "Demand-fixed", "No absorption — fixed customer base or budget-driven demand",
             "Government (fixed citizens), Education (fixed students)"],
        ],
        col_widths=[1.5, 3.0, 5.5, 6.0],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Scenario-stable: ")
    run.bold = True
    p.add_run("Demand elasticity is a structural property of the market, not a function of AI capability.")

    # -- T --
    doc.add_heading("T — Adoption Velocity", level=2)

    p = doc.add_paragraph()
    run = p.add_run("What it captures: ")
    run.bold = True
    p.add_run(
        "How fast the industry actually deploys AI, modeled as a logistic S-curve. This is the "
        "only variable that depends on both axes of the scenario framework — technology capability "
        "AND institutional friction."
    )

    eq6 = doc.add_paragraph()
    eq6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq6.paragraph_format.space_before = Pt(8)
    eq6.paragraph_format.space_after = Pt(8)
    run = eq6.add_run("T(t)  =  1 / (1 + exp(-1.2 × (t - t₀)))")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph(
        "The inflection point t₀ is derived from four sub-components, each scored 1–4 per "
        "sector and friction scenario:"
    )

    add_styled_table(doc,
        ["Sub-Component", "Direction", "What It Measures"],
        [
            ["T1 — Institutional Inertia", "Drag (higher = slower)", "Bureaucracy, change management, retraining costs, union resistance"],
            ["T2 — Systems Integration", "Drag (higher = slower)", "Legacy system complexity, data pipeline readiness, vendor ecosystem maturity"],
            ["T3 — Customer Acceptance", "Drag (higher = slower)", "End-user willingness to accept AI output, trust, brand risk"],
            ["T4 — Competitive Pressure", "Accelerant (higher = faster)", "Market dynamics pushing adoption — competitive intensity, margin pressure"],
        ],
        col_widths=[4.0, 3.5, 8.5],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The inflection point is computed as: D = avg(T1, T2, T3, T4), then t₀ = 1.5×D - 0.5×T4 - 0.5. "
        "Competitive pressure (T4) is effectively double-weighted — it appears in the average and "
        "directly reduces t₀ — reflecting the empirical observation that competitive dynamics are the "
        "strongest predictor of technology adoption speed."
    )

    p = doc.add_paragraph()
    run = p.add_run("Powder-keg flag: ")
    run.bold = True
    p.add_run(
        "When T < 0.2 at 18 months but (d_max × S(a) × E) > 0.10, the sector is flagged as a "
        "'powder keg' — high latent displacement potential with low current adoption. These sectors "
        "represent tail risk: if frictions suddenly ease, displacement could accelerate rapidly."
    )

    # -- R --
    doc.add_heading("R — Structural Resistance", level=2)

    p = doc.add_paragraph()
    run = p.add_run("What it captures: ")
    run.bold = True
    p.add_run(
        "Hard regulatory floors, licensure mandates, and liability regimes that cap displacement "
        "regardless of technical capability or market pressure."
    )

    doc.add_paragraph("R is derived from three sub-components, each scored 1–4:")

    add_styled_table(doc,
        ["Sub-Component", "What It Measures", "Example"],
        [
            ["f1 — Liability", "Who is liable when AI errs? Unresolved liability allocation blocks deployment",
             "Medical malpractice, legal advice liability"],
            ["f2 — Statutory Mandate", "Hard licensure requirements or legal mandates for human involvement",
             "Licensed physicians, CPAs signing audits, notarization"],
            ["f3 — Labor & Gatekeeping", "Unions, professional associations, and credentialing bodies that "
             "control workforce entry", "Bar associations, medical boards, teaching certificates"],
        ],
        col_widths=[3.5, 6.5, 6.0],
    )

    doc.add_paragraph("")

    eq7 = doc.add_paragraph()
    eq7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq7.paragraph_format.space_before = Pt(8)
    eq7.paragraph_format.space_after = Pt(8)
    run = eq7.add_run("F = f1 + f2 + f3        R = 1 - 0.7 × (F - 3) / 9")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True

    doc.add_paragraph(
        "At the floor (F = 3, no barriers), R = 1.00 — regulation imposes no cap. At the ceiling "
        "(F = 12, maximum barriers), R = 0.30 — regulation blocks 70% of potential displacement. "
        "This is scenario-stable: regulatory regimes do not change meaningfully within 18 months."
    )

    doc.add_page_break()

    # =========================================================================
    # 7. THE 2×2 SCENARIO FRAMEWORK
    # =========================================================================
    doc.add_heading("7.  The 2×2 Scenario Framework", level=1)

    doc.add_paragraph(
        "The model produces four displacement estimates per sector by crossing two independent "
        "axes of uncertainty. This avoids the false precision of a single-point forecast and helps "
        "clients plan across a range of plausible futures."
    )

    add_styled_table(doc,
        ["Scenario", "Tech Capability", "Adoption Friction", "Character"],
        [
            ["1", "Moderate", "Low", "Baseline — steady AI improvement, easy deployment"],
            ["2", "Moderate", "High", "Moderate tech, institutional resistance slows adoption"],
            ["3", "Significant", "Low", "Breakthrough AI, organizations deploy rapidly"],
            ["4", "Significant", "High", "Breakthrough AI, but institutional barriers slow rollout"],
        ],
        col_widths=[2.0, 3.5, 3.5, 7.0],
    )

    doc.add_paragraph("")
    doc.add_heading("How Scenarios Map to Variables", level=2)

    add_styled_table(doc,
        ["Variable", "Tech Axis", "Friction Axis", "Scenario Behavior"],
        [
            ["a(i,s) — Automatability", "Yes", "No", "Higher under Significant than Moderate scenarios"],
            ["d_max(i) — Displacement ceiling", "No", "No", "Fixed across all four scenarios"],
            ["E(i) — Elasticity dampener", "No", "No", "Fixed across all four scenarios"],
            ["T(t; i, s) — Adoption velocity", "Yes", "Yes", "Varies on both axes — fastest in Scenario 3, slowest in Scenario 2"],
            ["R(i) — Structural resistance", "No", "No", "Fixed across all four scenarios"],
        ],
        col_widths=[4.5, 2.0, 2.5, 7.0],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "The scenario framework enforces a monotonicity constraint: Significant capability scores "
        "must always be ≥ Moderate scores. More advanced AI cannot make a task less automatable. "
        "Similarly, low-friction scenarios must show adoption ≥ high-friction scenarios."
    )

    doc.add_page_break()

    # =========================================================================
    # 8. VALIDATION & QUALITY ASSURANCE
    # =========================================================================
    doc.add_heading("8.  Validation & Quality Assurance", level=1)

    doc.add_paragraph(
        "The model incorporates multiple layers of quality assurance to ensure scores are "
        "consistent, calibrated, and defensible."
    )

    doc.add_heading("Programmatic Audit Checks", level=2)

    checks = [
        ("Distribution skew", "No single variable has >35% of scores concentrated at one anchor point. "
         "Ensures meaningful differentiation across occupations."),
        ("Cross-variable correlation", "Correlation between workflow separability and substitutability "
         "must be ≤ 0.80. High correlation would indicate halo effects in scoring."),
        ("SOC-group coherence", "Within 2-digit SOC occupation groups, standard deviation < 0.40. "
         "Similar jobs should receive similar scores."),
        ("Logical consistency", "Flags combinations like low substitutability + high workflow separability "
         "(tension between 'human IS the product' and 'tasks easily offloaded'), or high digital task share "
         "+ low scalability (digital work that supposedly doesn't scale)."),
    ]
    for name, desc in checks:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}:  ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Reliability Verification", level=2)

    doc.add_paragraph(
        "A stratified 10% sample of occupations (46 jobs, sampled proportionally across sectors) "
        "is independently re-scored. Agreement metrics ensure reproducibility:"
    )

    add_styled_table(doc,
        ["Metric", "Threshold", "What It Measures"],
        [
            ["Weighted Quadratic Kappa", "κ ≥ 0.75", "Agreement beyond chance, weighted by distance between scores"],
            ["Within-One-Step Agreement", "≥ 90%", "Re-scored value within one scale step of original"],
            ["Exact Agreement", "≥ 65%", "Re-scored value matches original exactly"],
        ],
        col_widths=[4.5, 3.0, 8.5],
    )

    doc.add_heading("Scenario Monotonicity", level=2)

    doc.add_paragraph(
        "All task and job scores are verified to satisfy: Significant ≥ Moderate for every occupation. "
        "Any violations are investigated and corrected. This is not just a data check — it is a "
        "logical constraint of the model. More advanced AI capability cannot reduce automation potential."
    )

    doc.add_page_break()

    # =========================================================================
    # 9. KEY DESIGN PRINCIPLES
    # =========================================================================
    doc.add_heading("9.  Key Design Principles", level=1)

    principles = [
        ("Multiplicative, not additive",
         "Every factor in the master equation acts as a gate. A single bottleneck — workflow "
         "inseparability, regulatory resistance, slow adoption — is sufficient to compress displacement "
         "toward zero. This prevents the model from generating implausible estimates driven by one "
         "dominant factor."),
        ("Bottom-up, not top-down",
         "Displacement estimates are built from 5,383 individual task scores, not from assumptions about "
         "industries or job categories. The macro forecast is an emergent property of micro-level assessments."),
        ("Scenarios, not predictions",
         "The model does not claim to know which future will materialize. It brackets the range of plausible "
         "outcomes across two independent axes of uncertainty, allowing clients to stress-test their workforce "
         "strategies."),
        ("Independent reasoning to prevent halo effects",
         "Each scoring variable uses forced independent reasoning blocks. Agents must justify workflow "
         "separability without reference to scalability, and vice versa. Construct-specific task features are "
         "curated per variable to avoid contamination."),
        ("Human-IS-the-product vs. human preference",
         "The model draws a sharp line between intrinsic human necessity (x_sub = 0, the human IS the product "
         "— e.g., therapy, pastoral care, live performance) and customer preference (T3 — customers aren't yet "
         "willing to accept AI, but this preference can erode). The first is structural; the second is temporal."),
        ("Historical grounding",
         "Displacement ceilings (d_max) are anchored to observed labor market dynamics, not theoretical maxima. "
         "The model cannot generate displacement rates faster than any industry has ever experienced, regardless "
         "of AI capability."),
        ("Conservative by design",
         "The multiplicative structure, sigmoid compression, historical ceilings, and friction gates all push "
         "estimates toward the conservative end. The model is calibrated to avoid false alarms while still "
         "surfacing genuine vulnerability."),
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}.  ")
        run.bold = True
        run.font.color.rgb = NAVY
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    # -- Footer note --
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Methodology Appendix —")
    run.font.color.rgb = LIGHT_GRAY
    run.font.size = Pt(9)
    run.italic = True

    # =========================================================================
    # SAVE
    # =========================================================================
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
